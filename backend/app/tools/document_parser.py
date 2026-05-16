"""
Document Parser - Extract text from PDF, Word, and text documents
Optimized for high concurrency, memory safety, and robust encoding detection.
"""
import os
import anyio
import mimetypes
from typing import Dict, List, Optional, Any
from pathlib import Path
from loguru import logger
from concurrent.futures import ThreadPoolExecutor

# 安全依赖库引入
try:
    import charset_normalizer
except ImportError:
    charset_normalizer = None


class DocumentParser:
    """Parse various document formats and extract text asynchronously or synchronously"""
    
    def __init__(self, max_text_length: int = 10 * 1024 * 1024):
        """
        Args:
            max_text_length: 内存安全阈值，单文件默认最大提取 10MB 文本，防止 OOM
        """
        self.max_text_length = max_text_length
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.txt': self._parse_txt,
            '.md': self._parse_txt,
            '.rst': self._parse_txt,
            '.csv': self._parse_txt
        }
    
    def parse_file(self, file_path: str | Path) -> Dict[str, Any]:
        """
        [同步接口] 解析单个文档并提取文本和元数据
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Target document file not found at: {path}")
            
        ext = path.suffix.lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")
        
        try:
            parse_func = self.supported_formats[ext]
            result = parse_func(path)
            
            # 安全地封装全局元数据
            result['metadata'] = {
                'file_name': path.name,
                'file_path': str(path.absolute()),
                'file_size': path.stat().st_size,
                'file_type': ext,
                'mime_type': mimetypes.guess_type(path)[0] or "application/octet-stream"
            }
            
            # 文本截断防护提示
            if len(result.get('text', '')) >= self.max_text_length:
                result['warning'] = result.get('warning', '') + " Text content truncated due to size limits."
                
            logger.info(f"Successfully parsed {path.name}: {len(result['text'])} characters extracted.")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse document pipeline {file_path}: {e}")
            raise RuntimeError(f"Parser pipeline fault: {str(e)}")

    async def parse_file_async(self, file_path: str | Path) -> Dict[str, Any]:
        """
        [异步接口] 将阻塞的 I/O 解析过程托管到 AnyIO 专属线程池中执行，防止假死 FastAPI 事件循环
        """
        return await anyio.to_thread.run_sync(self.parse_file, file_path)

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """高效抽取 PDF 文本 (优先使用 pdfplumber, 降级使用新版 pypdf)"""
        try:
            import pdfplumber
            
            text_pages = []
            page_metadata = []
            total_chars = 0
            
            with pdfplumber.open(path) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, 1):
                    # 内存安全闸
                    if total_chars >= self.max_text_length:
                        break
                        
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        cleaned_text = page_text.strip()
                        text_pages.append(cleaned_text)
                        total_chars += len(cleaned_text)
                        page_metadata.append({
                            'page_number': page_num,
                            'char_count': len(cleaned_text)
                        })
            
            return {
                'text': "\n\n".join(text_pages)[:self.max_text_length],
                'pages': page_metadata,
                'total_pages': total_pages
            }
            
        except ImportError:
            logger.warning("pdfplumber not installed, falling back to pypdf standard package.")
            return self._parse_pdf_pypdf(path)
        except Exception as e:
            logger.error(f"pdfplumber failed for {path.name}: {e}")
            return self._parse_pdf_pypdf(path)

    def _parse_pdf_pypdf(self, path: Path) -> Dict[str, Any]:
        """基于现代 pypdf 库的兜底方案 (注意: 旧版 PyPDF2 已弃用)"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(path)
            text_pages = []
            total_chars = 0
            
            for page in reader.pages:
                if total_chars >= self.max_text_length:
                    break
                page_text = page.extract_text() or ""
                if page_text.strip():
                    cleaned_text = page_text.strip()
                    text_pages.append(cleaned_text)
                    total_chars += len(cleaned_text)
            
            return {
                'text': "\n\n".join(text_pages)[:self.max_text_length],
                'total_pages': len(reader.pages),
                'parser_fallback': True
            }
        except Exception as e:
            logger.error(f"All PDF parsers collapsed for {path.name}: {e}")
            raise

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """抽取 Word 文档 (.docx) 的段落及表格数据"""
        try:
            from docx import Document
            
            doc = Document(path)
            elements = []
            total_chars = 0
            
            # 1. 抽取段落
            for para in doc.paragraphs:
                if total_chars >= self.max_text_length:
                    break
                text = para.text.strip()
                if text:
                    elements.append(text)
                    total_chars += len(text)
            
            # 2. 抽取表格并将其转换为结构化的文本块
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_text = []
                for row in table.rows:
                    if total_chars >= self.max_text_length:
                        break
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    merged_table = "\n".join(table_text)
                    elements.append(f"\n[Table Data]:\n{merged_table}")
                    total_chars += len(merged_table)
            
            return {
                'text': "\n\n".join(elements)[:self.max_text_length],
                'paragraph_count': len(doc.paragraphs),
                'table_count': table_count
            }
        except Exception as e:
            logger.error(f"DOCX native parser crashed for {path.name}: {e}")
            raise

    def _parse_doc(self, path: Path) -> Dict[str, Any]:
        """安全解析老版二进制 .doc 文件 (杜绝 shell=True 的命令注入风险)"""
        try:
            import subprocess
            
            # 使用列表传入参数，完全断绝 shell=True 的字符串命令注入隐患
            result = subprocess.run(
                ['antiword', '-w', '0', str(path.absolute())],
                capture_output=True,
                text=True,
                timeout=15
            )
            
            if result.returncode == 0:
                return {'text': result.stdout[:self.max_text_length]}
            else:
                raise RuntimeError(result.stderr or "antiword process error")
                
        except Exception as e:
            logger.warning(f"OS-level .doc parsing failed or antiword is not configured: {e}")
            return {
                'text': "",
                'warning': f"Limited support for legacy .doc format. Conversion to .docx is recommended. Error: {str(e)}"
            }

    def _parse_txt(self, path: Path) -> Dict[str, Any]:
        """基于字符集智能嗅探的高可靠纯文本读取机制"""
        # 1. 优先使用智能嗅探器（杜绝盲猜遍历导致的解码崩溃）
        if charset_normalizer:
            try:
                with open(path, 'rb') as f:
                    # 仅读取前 64KB 特征字节数据，兼顾超大文件的检测效率
                    raw_data = f.read(64 * 1024)
                matches = charset_normalizer.from_bytes(raw_data)
                best_match = matches.best()
                
                if best_match and best_match.encoding:
                    encoding = best_match.encoding
                    with open(path, 'r', encoding=encoding, errors='replace') as f:
                        return {
                            'text': f.read(self.max_text_length),
                            'encoding': encoding,
                            'detection_method': 'charset_normalizer'
                        }
            except Exception as e:
                logger.debug(f"Smart normalizer failed, switching to backup loops: {e}")

        # 2. 兜底策略：标准行业编码序列循环重试机制
        fallback_encodings = ['utf-8', 'gbk', 'utf-8-sig', 'gb18030', 'latin-1']
        for encoding in fallback_encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return {
                        'text': f.read(self.max_text_length),
                        'encoding': encoding,
                        'detection_method': 'fallback_loop'
                    }
            except (UnicodeDecodeError, LookupError):
                continue
                
        # 3. 终极自适应：若全部失败，带损坏标示强制读取，确保系统可用性
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            return {
                'text': f.read(self.max_text_length),
                'encoding': 'utf-8_corrupted_fallback',
                'warning': 'Character decoding was partially corrupted; unreadable entities replaced.'
            }

    def batch_parse(self, file_paths: List[str | Path], max_workers: int = 4) -> List[Dict[str, Any]]:
        """
        [同步多线程加速] 批量并发并行调度解析文档
        """
        results = []
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_path = {executor.submit(self.parse_file, path): path for path in file_paths}
            for future in future_to_path:
                path = future_to_path[future]
                try:
                    results.append(future.result())
                except Exception as e:
                    results.append({
                        'file_path': str(path),
                        'text': "",
                        'error': f"Batch processing thread faulted: {str(e)}"
                    })
        return results

    async def batch_parse_async(self, file_paths: List[str | Path]) -> List[Dict[str, Any]]:
        """
        [异步并发加速] 利用 AnyIO TaskGroup 实现高吞吐量异步非阻塞并行文件解析
        """
        results = []
        
        async def _worker(path: str | Path):
            try:
                res = await self.parse_file_async(path)
                results.append(res)
            except Exception as e:
                results.append({
                    'file_path': str(path),
                    'text': "",
                    'error': f"Async worker collapsed: {str(e)}"
                })

        async with anyio.create_task_group() as tg:
            for file_path in file_paths:
                tg.start_soon(_worker, file_path)
                
        return results


# 便利的快捷顶层全局调用接口（非阻塞异步设计，完美对接 FastAPI）
async def parse_document(file_path: str | Path) -> Dict[str, Any]:
    """
    快速全自动异步解析外部文档
    """
    parser = DocumentParser()
    return await parser.parse_file_async(file_path)