"""
Test Document Parser
"""
import os
import sys
from pathlib import Path

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent))

from app.tools.document_parser import DocumentParser


def test_parser():
    """Test document parser with sample files"""
    parser = DocumentParser()
    
    # Create test directory
    test_dir = Path("test_documents")
    test_dir.mkdir(exist_ok=True)
    
    print("=" * 60)
    print("Document Parser Test")
    print("=" * 60)
    
    # Test 1: Create and parse a simple text file
    print("\n1. Testing TXT file...")
    txt_file = test_dir / "test.txt"
    txt_file.write_text("这是测试文本内容。\nHello, this is a test.", encoding='utf-8')
    
    result = parser.parse_file(str(txt_file))
    print(f"[OK] TXT parsed: {len(result['text'])} characters")
    print(f"  Content preview: {result['text'][:50]}...")
    
    # Test 2: Create and parse a Word document
    print("\n2. Testing DOCX file...")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个 Word 文档测试。')
        doc.add_paragraph('This is a test paragraph in Chinese and English.')
        doc.add_heading('第二章', level=1)
        doc.add_paragraph('更多内容...')
        
        docx_file = test_dir / "test.docx"
        doc.save(str(docx_file))
        
        result = parser.parse_file(str(docx_file))
        print(f"[OK] DOCX parsed: {len(result['text'])} characters")
        print(f"  Paragraphs: {result.get('paragraph_count', 'N/A')}")
        print(f"  Content preview: {result['text'][:100]}...")
    except Exception as e:
        print(f"[FAIL] DOCX test failed: {e}")
    
    # Test 3: Try to parse a PDF (if available)
    print("\n3. Testing PDF file...")
    # Check if there's any PDF in the system for testing
    sample_pdf = None
    for pdf_path in Path(".").rglob("*.pdf"):
        if pdf_path.stat().st_size < 10 * 1024 * 1024:  # < 10MB
            sample_pdf = pdf_path
            break
    
    if sample_pdf:
        try:
            result = parser.parse_file(str(sample_pdf))
            print(f"[OK] PDF parsed: {len(result['text'])} characters")
            print(f"  Pages: {result.get('total_pages', 'N/A')}")
            print(f"  Content preview: {result['text'][:100]}...")
        except Exception as e:
            print(f"[FAIL] PDF test failed: {e}")
    else:
        print("[SKIP] No PDF file available for testing")
        print("  To test PDF parsing, place a PDF file in the directory")
    
    # Cleanup
    print("\n" + "=" * 60)
    print("Test completed!")
    print(f"Test files created in: {test_dir.absolute()}")
    print("=" * 60)


if __name__ == "__main__":
    test_parser()
